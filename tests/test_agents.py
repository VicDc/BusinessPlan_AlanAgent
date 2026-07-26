import json
import pytest
from pathlib import Path
import os
from unittest.mock import AsyncMock, MagicMock, patch

from app.core.types import (
    BusinessIdeaProfile, FounderProfile, ProductType, AgentOutput,
    RevisionStatus, ChartSpec
)
from app.agents.intake_agent import IntakeAgent
from app.agents.vision_agent import VisionAgent
from app.agents.market_agent import MarketAgent
from app.agents.team_agent import TeamAgent
from app.agents.setup_agent import SetupAgent
from app.agents.financial_agent import FinancialAgent
from app.agents.funding_agent import FundingAgent
from app.agents.orchestrator import Orchestrator
from app.services.charts import render_chart_specs
from app.services.report_builder import markdown_to_docx, render_agent_section, save_markdown_report
from docx import Document
from app.services.llm_logging import new_call_id


@pytest.fixture(autouse=True)
def mock_llm_logging_path(tmp_path, monkeypatch):
    import app.services.llm_logging
    monkeypatch.setattr(app.services.llm_logging, "LOGS_DIR", tmp_path)
    monkeypatch.setattr(app.services.llm_logging, "LOGS_FILE", tmp_path / "test_llm_calls.jsonl")
    if hasattr(app.services.llm_logging, "LOG_FILE"):
        monkeypatch.setattr(app.services.llm_logging, "LOG_FILE", tmp_path / "test_llm_calls.jsonl")


class MockLLMService:
    def __init__(self, responses=None):
        self.responses = responses or {}
        self.generate_calls = []
        self.run_id = None
        self.iteration = None
        self.last_call_id = None

    async def generate(self, system_prompt, user_message, temperature=0.2, max_tokens=2000, **kwargs):
        self.last_call_id = new_call_id()
        self.generate_calls.append((system_prompt, user_message))
        # Match based on unique strings in the prompt content
        if "trasforma le risposte" in system_prompt:
            return self.responses.get("intake", "{}")
        elif "ideazione imprenditoriale" in system_prompt:
            return self.responses.get("vision", "{}")
        elif "analista di mercato" in system_prompt:
            return self.responses.get("market", "{}")
        elif "organizzazione di team" in system_prompt:
            return self.responses.get("team", "{}")
        elif "adempimenti legali" in system_prompt:
            return self.responses.get("setup", "{}")
        elif "modellazione finanziaria" in system_prompt:
            return self.responses.get("financial", "{}")
        elif "finanziamenti per startup" in system_prompt:
            return self.responses.get("funding", "{}")
        elif "partner di un team" in system_prompt:
            return self.responses.get("orchestrator", "{}")
        elif "redattore finale" in system_prompt:
            return self.responses.get("report_writer", "")
        return "{}"


@pytest.fixture
def base_profile():
    return BusinessIdeaProfile(
        project_name="TestProject",
        idea_description="Test Description",
        need_addressed="Test Need",
        product_type=ProductType.SERVICE,
        sector_hint="F&B",
        target_region="Italy",
        founders=[
            FounderProfile(name="Alice", skills=["cooking"], availability="full-time")
        ],
        available_capital_eur=5000.0,
        desired_timeline_months=12,
        notes="None"
    )


@pytest.mark.asyncio
async def test_intake_agent():
    mock_response = {
        "project_name": "Fermenta",
        "idea_description": "Produzione kombucha",
        "need_addressed": "Mancanza drink analcolici",
        "product_type": "physical",
        "sector_hint": "F&B",
        "target_region": "Milano",
        "founders": [{"name": "Bob", "skills": ["kombucha"], "availability": "part-time"}],
        "available_capital_eur": 3000.0,
        "desired_timeline_months": 6,
        "notes": "",
        "raw_section_notes": {"vision": "kombucha vision"},
        "needs_clarification": ["Mancano costi"],
        "summary_markdown": "# Fermenta\nSintesi",
        "confidence": 0.8
    }
    llm = MockLLMService({"intake": json.dumps(mock_response)})
    intake = IntakeAgent(llm)
    
    # Test template generation
    template = intake.generate_template_markdown()
    assert "# Business Plan — Modulo di Intake" in template
    
    # Test parsing brief
    report = await intake.parse_brief("Raw user responses markdown")
    assert report.profile.project_name == "Fermenta"
    assert report.profile.product_type == ProductType.PHYSICAL
    assert report.profile.founders[0].name == "Bob"
    assert report.needs_clarification == ["Mancano costi"]
    assert report.confidence == 0.8


@pytest.mark.asyncio
async def test_specialist_agents_success(base_profile):
    mock_responses = {
        "vision": json.dumps({"confidence": 0.9, "value_proposition": "Val prop"}),
        "market": json.dumps({"confidence": 0.8, "market_sizing": {"som_eur": "50k"}}),
        "team": json.dumps({"confidence": 0.85, "skills_gap": []}),
        "setup": json.dumps({"confidence": 0.7, "legal_form": {"recommended": "SRL"}}),
        "financial": json.dumps({"confidence": 0.8, "pricing": {"unit_price_eur": 10.0}}),
        "funding": json.dumps({"confidence": 0.9, "funding_gap_eur": 1000.0})
    }
    llm = MockLLMService(mock_responses)
    
    vision = VisionAgent(llm)
    v_out = await vision.process(base_profile)
    assert v_out.status == "success"
    assert v_out.data["value_proposition"] == "Val prop"
    
    market = MarketAgent(llm)
    m_out = await market.process(base_profile)
    assert m_out.status == "success"
    assert m_out.data["market_sizing"]["som_eur"] == "50k"
    
    team = TeamAgent(llm)
    t_out = await team.process(base_profile)
    assert t_out.status == "success"
    
    setup = SetupAgent(llm)
    s_out = await setup.process(base_profile)
    assert s_out.status == "success"
    
    financial = FinancialAgent(llm)
    f_out = await financial.process(base_profile)
    assert f_out.status == "success"
    assert f_out.data["pricing"]["unit_price_eur"] == 10.0
    
    funding = FundingAgent(llm)
    fu_out = await funding.process(base_profile)
    assert fu_out.status == "success"


@pytest.mark.asyncio
async def test_specialist_agents_json_decode_error(base_profile):
    llm = MockLLMService({"vision": "Not a valid JSON"})
    vision = VisionAgent(llm)
    v_out = await vision.process(base_profile)
    assert v_out.status == "error"
    assert v_out.confidence == 0.0
    assert "Parsing JSON fallito" in v_out.reasoning


@pytest.mark.asyncio
async def test_market_and_funding_web_search(base_profile):
    mock_search = AsyncMock()
    mock_search.search.return_value = [{"title": "Search Title", "snippet": "Search Snippet", "link": "http://link.com"}]
    
    llm = MockLLMService({
        "market": json.dumps({"confidence": 0.8, "sources_used": ["http://link.com"]}),
        "funding": json.dumps({"confidence": 0.8, "funding_sources": []})
    })
    
    market = MarketAgent(llm, web_search_service=mock_search)
    await market.process(base_profile)
    mock_search.search.assert_called_with("mercato F&B Italy", max_results=3)
    
    funding = FundingAgent(llm, web_search_service=mock_search)
    await funding.process(base_profile)
    mock_search.search.assert_called_with("bandi agevolazioni startup F&B Italy", max_results=3)


@pytest.mark.asyncio
async def test_orchestrator_approved_flow(base_profile, tmp_path):
    mock_responses = {
        "vision": json.dumps({"confidence": 0.9}),
        "market": json.dumps({"confidence": 0.9}),
        "team": json.dumps({"confidence": 0.9}),
        "setup": json.dumps({"confidence": 0.9}),
        "financial": json.dumps({
            "confidence": 0.9,
            "charts_needed": [
                {
                    "chart_type": "bar",
                    "title": "Proiezioni",
                    "labels": ["Anno 1"],
                    "series": {"Ricavi": [100.0]},
                    "filename": "test_chart"
                }
            ]
        }),
        "funding": json.dumps({"confidence": 0.9}),
        "orchestrator": json.dumps({
            "status": "APPROVED",
            "revisions_needed": [],
            "confidence_overall": 0.9,
            "iteration": 1
        }),
        "report_writer": "# Business Plan Test\nFine."
    }
    
    llm = MockLLMService(mock_responses)
    orchestrator = Orchestrator(llm)
    
    with patch("app.services.charts.CHARTS_DIR", tmp_path), \
         patch("app.services.report_builder.OUTPUT_DIR", tmp_path), \
         patch("docx.document.Document.add_picture") as mock_add_picture:
        result = await orchestrator.run(base_profile)
        
        assert result.status == RevisionStatus.APPROVED
        assert result.total_iterations == 1
        assert "Business Plan Test" in result.business_plan_markdown
        # Verify chart generated
        assert len(result.charts_generated) == 1
        assert Path(result.charts_generated[0]).exists()
        # Verify DOCX built
        assert Path(result.business_plan_docx_path).exists()
        # Verify MD built
        assert Path(result.business_plan_md_path).exists()


@pytest.mark.asyncio
async def test_orchestrator_revision_loop_and_max_cycles(base_profile, tmp_path):
    orch_responses = [
        json.dumps({
            "status": "REVISION_NEEDED",
            "revisions_needed": [
                {"agent": "vision", "issue": "Val prop missing", "correction_context": "Add val prop"}
            ],
            "business_plan_markdown": "",
            "confidence_overall": 0.5,
            "iteration": 1
        }),
        json.dumps({
            "status": "APPROVED",
            "revisions_needed": [],
            "business_plan_markdown": "# Plan final approved",
            "confidence_overall": 0.9,
            "iteration": 2
        })
    ]
    
    class MultiOrchLLMService(MockLLMService):
        def __init__(self, responses):
            super().__init__(responses)
            self.orch_call_count = 0
            
        async def generate(self, system_prompt, user_message, temperature=0.2, max_tokens=2000, **kwargs):
            self.last_call_id = new_call_id()
            if "partner di un team" in system_prompt:
                resp = orch_responses[self.orch_call_count]
                self.orch_call_count += 1
                return resp
            return await super().generate(system_prompt, user_message, temperature, max_tokens, **kwargs)

    mock_responses = {
        "vision": json.dumps({"confidence": 0.9}),
        "market": json.dumps({"confidence": 0.9}),
        "team": json.dumps({"confidence": 0.9}),
        "setup": json.dumps({"confidence": 0.9}),
        "financial": json.dumps({"confidence": 0.9}),
        "funding": json.dumps({"confidence": 0.9}),
    }
    
    llm = MultiOrchLLMService(mock_responses)
    orchestrator = Orchestrator(llm)
    
    # Ermetico: lo scenario richiede >=2 iterazioni, non dipendere dal
    # MAX_REVISION_CYCLES ambientale (.env può impostarlo a 1).
    with patch("app.agents.orchestrator.settings.MAX_REVISION_CYCLES", 3), \
         patch("app.services.charts.render_chart_specs", return_value=[]), \
         patch("app.agents.orchestrator.markdown_to_docx", return_value="dummy_path.docx"), \
         patch("app.agents.orchestrator.save_markdown_report", return_value="dummy_path.md"), \
         patch("app.services.report_builder.OUTPUT_DIR", tmp_path):
        result = await orchestrator.run(base_profile)

        assert result.status == RevisionStatus.APPROVED
        assert result.total_iterations == 2
        assert len(result.revision_log) == 1
        assert result.revision_log[0]["revisions"][0]["agent"] == "vision"
        assert result.business_plan_docx_path == "dummy_path.docx"
        assert result.business_plan_md_path == "dummy_path.md"


def test_charts_rendering_valid_and_invalid(tmp_path):
    with patch("app.services.charts.CHARTS_DIR", tmp_path):
        specs = [
            # Valid spec
            {
                "chart_type": "bar",
                "title": "Bar Title",
                "labels": ["A", "B"],
                "series": {"S1": [1.0, 2.0]},
                "filename": "chart1"
            },
            # Invalid spec
            {
                "chart_type": "invalid_type",
                "title": "Error Title",
                "labels": [],
                "series": {},
                "filename": "chart2"
            }
        ]
        paths = render_chart_specs(specs)
        
        # Valid chart is rendered
        assert len(paths) == 1
        assert Path(paths[0]).name == "chart1.png"
        assert Path(paths[0]).exists()


def test_report_builder_docx(tmp_path):
    with patch("app.services.report_builder.OUTPUT_DIR", tmp_path), \
         patch("docx.document.Document.add_picture") as mock_add_picture:
        markdown_text = """# Business Plan
## 1. Executive Summary
This is a test summary.
- Bullet point 1
- Bullet point 2
1. Number point 1
2. Number point 2
"""
        dummy_chart = tmp_path / "dummy_chart.png"
        dummy_chart.write_text("dummy")
        
        docx_path = markdown_to_docx(
            markdown_text=markdown_text,
            chart_paths=[str(dummy_chart)],
            output_filename="test_report.docx"
        )
        
        assert Path(docx_path).exists()
        assert Path(docx_path).name == "test_report.docx"
        mock_add_picture.assert_called_once()


def test_report_builder_markdown(tmp_path):
    with patch("app.services.report_builder.OUTPUT_DIR", tmp_path):
        markdown_text = """# Business Plan
## 1. Executive Summary
This is a test summary.
"""
        dummy_chart = tmp_path / "dummy_chart.png"
        dummy_chart.write_text("dummy")
        
        md_path = save_markdown_report(
            markdown_text=markdown_text,
            chart_paths=[str(dummy_chart)],
            output_filename="test_report.md"
        )
        
        assert Path(md_path).exists()
        assert Path(md_path).name == "test_report.md"
        content = Path(md_path).read_text(encoding="utf-8")
        assert "Business Plan" in content
        assert "dummy_chart.png" in content


def test_render_agent_section_is_readable_not_json():
    data = {
        "value_proposition": "Pane fresco a domicilio",
        "need_validation": {"is_clear": True, "comment": "Bisogno concreto"},
        "competitors": [
            {"name": "Forno A", "strengths": "vicino", "weaknesses": "caro"},
            {"name": "Forno B", "strengths": "economico", "weaknesses": "lontano"},
        ],
        "risk_flags": ["mercato piccolo", "stagionalità"],
        "confidence": 0.8,          # deve essere omesso
        "charts_needed": [{"x": 1}],  # deve essere omesso
    }
    md = render_agent_section("vision", data)

    assert "```json" not in md
    assert "## Vision" in md
    assert "**Value Proposition:** Pane fresco a domicilio" in md
    assert "| Name | Strengths | Weaknesses |" in md   # lista di dict → tabella
    assert "- mercato piccolo" in md                    # lista di stringhe → bullet
    assert "confidence" not in md.lower()               # chiave tecnica omessa
    assert "charts_needed" not in md


def test_markdown_to_docx_tables_and_bold(tmp_path):
    md = (
        "#### Sotto-sezione\n"
        "\n"
        "Un paragrafo con **testo importante** in grassetto.\n"
        "\n"
        "| Nome | Ruolo |\n"
        "| --- | --- |\n"
        "| **Anna** | CEO |\n"
        "| Marco | CTO |\n"
        "\n"
        "⚠️ Avviso di prova.\n"
    )
    with patch("app.services.report_builder.OUTPUT_DIR", tmp_path):
        out = markdown_to_docx(md, chart_paths=[], output_filename="t.docx")

    doc = Document(out)

    # 1. almeno una tabella vera
    assert len(doc.tables) > 0
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Nome"      # header, separatore ignorato
    assert table.rows[1].cells[0].text == "Anna"      # ** rimossi anche in cella

    # 2. run in grassetto presente nei paragrafi
    bold_runs = [r.text for p in doc.paragraphs for r in p.runs if r.bold and r.text.strip()]
    assert "testo importante" in bold_runs

    # 3. nessun asterisco o pipe letterale residuo nei paragrafi o nelle celle
    para_text = "\n".join(p.text for p in doc.paragraphs)
    cell_text = "\n".join(c.text for row in table.rows for c in row.cells)
    assert "*" not in para_text and "|" not in para_text
    assert "*" not in cell_text and "|" not in cell_text
