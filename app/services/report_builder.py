"""
Converte il business plan Markdown finale in DOCX, incorporando i grafici
già renderizzati da charts.py. Riuso adattato di
src/tools/docx_writer.py da strategic-consulting-crew.
"""
from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "output")) / "reports"

# Chiavi tecniche mai utili in un report leggibile
_OMIT_KEYS = {"confidence", "charts_needed"}


def _titleize(key: str) -> str:
    return str(key).replace("_", " ").strip().title()


def _scalar_to_str(v) -> str:
    """Appiattisce un valore (anche list/dict annidati) in una stringa da cella
    tabella o bullet."""
    if isinstance(v, bool):
        return "Sì" if v else "No"
    if isinstance(v, list):
        return ", ".join(_scalar_to_str(x) for x in v)
    if isinstance(v, dict):
        return "; ".join(f"{_titleize(k)}: {_scalar_to_str(x)}" for k, x in v.items())
    return str(v)


def _render_table(rows: list[dict]) -> list[str]:
    cols = []
    for r in rows:
        for k in r:
            if k not in cols:
                cols.append(k)
    out = [
        "| " + " | ".join(_titleize(c) for c in cols) + " |",
        "| " + " | ".join("---" for _ in cols) + " |",
    ]
    for r in rows:
        out.append("| " + " | ".join(_scalar_to_str(r.get(c, "")) for c in cols) + " |")
    return out


def _render_value(key: str, value, level: int, out: list[str]) -> None:
    label = _titleize(key)
    heading = "#" * min(level, 6)
    if isinstance(value, dict):
        out.append(f"\n{heading} {label}\n")
        _render_dict(value, level + 1, out)
    elif isinstance(value, list) and value and all(isinstance(i, dict) for i in value):
        out.append(f"\n{heading} {label}\n")
        keysets = [list(d.keys()) for d in value]
        if all(ks == keysets[0] for ks in keysets):
            out.extend(_render_table(value))          # chiavi omogenee → tabella
        else:
            for item in value:                        # eterogenee → bullet appiattiti
                out.append(f"- {_scalar_to_str(item)}")
    elif isinstance(value, list):
        out.append(f"\n{heading} {label}\n")
        for item in value:
            out.append(f"- {_scalar_to_str(item)}")
    else:
        out.append(f"**{label}:** {_scalar_to_str(value)}")


def _render_dict(data: dict, level: int, out: list[str]) -> None:
    for key, value in data.items():
        if key in _OMIT_KEYS:
            continue
        _render_value(key, value, level, out)


def render_agent_section(agent_key: str, data: dict) -> str:
    """Converte il dict di output di un agente in markdown human-readable:
    - chiavi snake_case → titoli ### in Title Case leggibile
    - liste di dict con chiavi omogenee → tabelle markdown
    - liste di stringhe → bullet list
    - valori scalari → riga "**Chiave:** valore"
    - dict annidati → sottosezioni
    - ometti sempre le chiavi tecniche: "confidence", "charts_needed"
    Generico, guidato dalla struttura del dict, senza mappature hardcoded."""
    out = [f"## {_titleize(agent_key)}", ""]
    if isinstance(data, dict):
        _render_dict(data, 3, out)
    else:
        out.append(_scalar_to_str(data))
    return "\n".join(out)


def build_draft_markdown(profile, agent_outputs: dict, iteration: int, issues: list[str]) -> str:
    """Business plan PARZIALE con l'ultimo output di ciascun agente e i
    problemi segnalati in questa iterazione."""
    lines = [
        f"# Bozza — Iterazione {iteration}\n",
        "⚠️ Non ancora approvata dall'Orchestrator.\n",
    ]
    for agent_key, ao in agent_outputs.items():
        lines.append("")
        lines.append(render_agent_section(agent_key, ao.data))
    if issues:
        lines.append("\n## Problemi segnalati in questa iterazione\n")
        for issue in issues:
            if issue:
                lines.append(f"- {issue}")
    return "\n".join(lines)


def _add_runs_with_bold(paragraph, text: str, force_bold: bool = False) -> None:
    """Aggiunge `text` a un paragrafo interpretando **...** come run in
    grassetto. force_bold rende grassetto anche il testo non marcato (celle
    di intestazione tabella)."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            paragraph.add_run(part[2:-2]).bold = True
        else:
            run = paragraph.add_run(part)
            if force_bold:
                run.bold = True


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and len(line) >= 2


def _is_separator_row(line: str) -> bool:
    """Riga separatrice markdown: solo |, -, : e spazi."""
    return set(line) <= set("|-: ") and "-" in line


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, block: list[str]) -> None:
    rows = [_split_row(l) for l in block if not _is_separator_row(l)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""  # svuota il paragrafo di default
            _add_runs_with_bold(cell.paragraphs[0], row[ci] if ci < len(row) else "",
                                force_bold=(ri == 0))


def markdown_to_docx(
    markdown_text: str,
    chart_paths: list[str],
    output_filename: str = "business_plan.docx"
) -> str:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if _is_table_row(line):
            block = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                block.append(lines[i].strip())
                i += 1
            _add_table(doc, block)
            continue

        if not line:
            doc.add_paragraph("")
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            _add_runs_with_bold(doc.add_paragraph(style="List Bullet"), line[2:])
        elif re.match(r"^\d+\. ", line):
            _add_runs_with_bold(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", line))
        else:
            _add_runs_with_bold(doc.add_paragraph(), line)
        i += 1

    # Incorpora tutti i grafici renderizzati in coda al documento, in una
    # sezione dedicata (semplice e robusto; l'inserimento posizionale nel
    # testo può essere raffinato in una versione successiva)
    if chart_paths:
        doc.add_heading("Grafici e Proiezioni", level=1)
        for path in chart_paths:
            if Path(path).exists():
                doc.add_picture(path, width=Inches(5.5))

    output_path = OUTPUT_DIR / output_filename
    doc.save(str(output_path))
    return str(output_path)


def save_markdown_report(
    markdown_text: str,
    chart_paths: list[str],
    output_filename: str = "business_plan.md"
) -> str:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / output_filename

    content = markdown_text
    if chart_paths:
        content += "\n\n# Grafici e Proiezioni\n"
        for path in chart_paths:
            p = Path(path)
            content += f"\n![{p.stem}](file:///{p.as_posix()})\n"

    output_path.write_text(content, encoding="utf-8")
    return str(output_path)
